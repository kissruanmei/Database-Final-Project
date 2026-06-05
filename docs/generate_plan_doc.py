"""生成项目计划书 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- 全局样式 ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

# ---- 辅助函数 ----
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '黑体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('校园机动车综合管理平台 — 项目计划书')
run.font.size = Pt(18)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()  # 空行

# ========== 一、设计实现步骤 ==========
add_heading('一、数据库应用系统设计实现步骤', level=1)

make_table(
    ['步骤', '基本任务', '涉及知识 / 方法 / 工具'],
    [
        ['1. 需求分析',   '梳理业务流程，画 DFD，编写数据字典', '需求调研、DFD 图、数据字典'],
        ['2. 概念设计',   '识别实体/属性/联系，画 E-R 图',      'E-R 模型、draw.io'],
        ['3. 逻辑设计',   'E-R → 关系模式，确定主外键，规范化到 3NF', '关系模型、范式理论'],
        ['4. 物理设计',   'SQL Server 建库建表，加约束和索引',   'SSMS、T-SQL DDL、索引策略'],
        ['5. 实施与开发', '导入数据、编写视图/存储过程/触发器、开发 CLI', 'T-SQL DML、Python + pyodbc'],
        ['6. 测试与验收', '功能测试、数据完整性测试、答辩准备',   '测试用例、PPT / 文档'],
    ],
    col_widths=[3, 7, 6]
)

# ========== 二、已有基础 & 问题清单 ==========
add_heading('二、已有基础与问题清单', level=1)

add_heading('2.1 已掌握内容', level=2)
make_table(
    ['知识点', '状态'],
    [
        ['SQL 基础（SELECT / INSERT / UPDATE / DELETE）', '已掌握'],
        ['表的创建与约束（PK / FK / CHECK / UNIQUE）',     '已掌握'],
        ['E-R 图基本概念',                                 '已掌握'],
        ['范式理论（1NF / 2NF / 3NF）',                   '已掌握'],
        ['视图（View）',                                   '未学，需自学'],
        ['存储过程（Stored Procedure）',                   '未学，需自学'],
        ['触发器（Trigger）',                              '未学，需自学'],
        ['Python 连接 SQL Server（pyodbc）',               '未学，需自学'],
    ],
    col_widths=[10, 5]
)

add_heading('2.2 问题清单', level=2)
make_table(
    ['编号', '问题', '状态'],
    [
        ['1', '存储过程、视图、触发器未学过',          '需自学'],
        ['2', 'pyodbc 从未用过',                       '需学习'],
        ['3', 'DFD 数据流图没画过',                    '已解决'],
        ['4', 'E-R 图规范画法不熟练',                  '已解决'],
        ['5', '测试数据如何构造有代表性',              '已解决'],
    ],
    col_widths=[1.5, 9, 4]
)

add_heading('2.3 主要难点', level=2)
make_table(
    ['难点', '说明'],
    [
        ['难点 1（最大）', '存储过程/触发器从零自学并在项目中实际应用'],
        ['难点 2',         '3NF 规范化 — 消除冗余且不丢失业务语义'],
        ['难点 3',         '多表关联查询的性能与正确性'],
        ['难点 4',         'Python CLI 与 SQL Server 的稳定连接和异常处理'],
    ],
    col_widths=[3.5, 12]
)

# ========== 三、拓展需求 ==========
add_heading('三、拓展部分需求（初步）', level=1)

make_table(
    ['拓展方向', '具体内容', '难度'],
    [
        ['视图',     '常用多表查询封装（车主违规历史、本月违规统计）', '★☆☆'],
        ['存储过程', '违规处理流程（自动记录、累计积分、触发处罚）',   '★★☆'],
        ['触发器',   '违规插入时自动更新积分；通行证到期自动失效',     '★★☆'],
        ['统计报表', '按月/学院/违规类型统计数量',                     '★☆☆'],
        ['权限管理', '管理员/车主不同操作权限',                        '★★☆'],
    ],
    col_widths=[2.5, 10, 2.5]
)

p = doc.add_paragraph()
run = p.add_run('建议选取：视图 + 存储过程 + 触发器（最能体现课程知识点，答辩加分明显）')
run.bold = True
run.font.size = Pt(10)

# ========== 四、时间进度表 ==========
add_heading('四、时间进度表', level=1)

make_table(
    ['周次', '时间', '任务', '产出'],
    [
        ['第1周', '第1-2天',  '需求分析：梳理业务流程，画 DFD，写数据字典', '需求文档'],
        ['',      '第3-4天',  '概念设计：画 E-R 图',                        'E-R 图'],
        ['',      '第5-6天',  '逻辑设计：E-R → 表结构，规范化检查',         '表结构文档'],
        ['',      '第7天',    '环境搭建：安装 SQL Server + Python + pyodbc', '开发环境就绪'],
        ['第2周', '第8-9天',  '物理设计：建库建表，加约束和索引',            'DDL 脚本'],
        ['',      '第10天',   '插入测试数据 + 自学视图',                    '测试数据 + 视图'],
        ['',      '第11-12天','自学存储过程 + 编写存储过程',                '存储过程脚本'],
        ['',      '第13天',   'Python CLI 界面开发',                        'CLI 程序'],
        ['',      '第14天',   '集成测试：CLI 调用数据库验证功能',            '测试通过'],
        ['第3周', '第15天',   '自学触发器 + 编写触发器',                    '触发器脚本'],
        ['',      '第16天',   '补充拓展功能',                               '拓展功能完成'],
        ['',      '第17-18天','编写课程设计报告',                           '报告初稿'],
        ['',      '第19-20天','答辩准备：PPT + 演示排练',                   '答辩材料'],
        ['',      '第21天',   '预留缓冲 / 修改完善',                        '最终版本'],
    ],
    col_widths=[1.5, 2.5, 8, 4]
)

# ========== 五、组员分工 ==========
add_heading('五、组员分工', level=1)

make_table(
    ['角色', '人员', '职责'],
    [
        ['A — 数据库设计与开发', '许敬淞', '需求分析、E-R 图、表结构设计、DDL 建表、视图、存储过程、触发器、测试数据'],
        ['B — 应用开发与文档',   '郑奎昊', 'Python CLI 开发、数据库连接层、联调、课程设计报告、答辩 PPT'],
    ],
    col_widths=[4, 2, 10]
)

add_heading('协作点', level=2)
for item in ['需求分析阶段一起讨论业务流程', 'E-R 图一起审核确认', '联调测试阶段一起验证', '答辩演示一起准备']:
    doc.add_paragraph(item, style='List Bullet')

# ========== 六、核心技术栈 ==========
add_heading('六、核心技术栈', level=1)

make_table(
    ['组件', '技术选型'],
    [
        ['数据库',       'SQL Server（课程指定）'],
        ['管理工具',     'SQL Server Management Studio (SSMS)'],
        ['编程语言',     'Python 3.x'],
        ['数据库连接',   'pyodbc'],
        ['用户界面',     '命令行 CLI（input/print 菜单交互）'],
        ['文档工具',     'Word + PowerPoint'],
        ['画图工具',     'draw.io'],
    ],
    col_widths=[3, 13]
)

# ---- 保存 ----
doc.save('E:/DataBase-Homework/docs/项目计划书.docx')
print('文档已生成: docs/项目计划书.docx')
