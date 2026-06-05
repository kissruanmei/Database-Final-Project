# -*- coding: utf-8 -*-
"""
生成第二周任务完成报告 Word 文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def create_week2_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)

    # ===== 标题 =====
    title = doc.add_heading('校园机动车综合管理平台', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('第二周任务完成报告', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 基本信息表 =====
    doc.add_heading('一、基本信息', level=2)

    info_table = doc.add_table(rows=5, cols=4, style='Table Grid')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ['项目名称', '校园机动车综合管理平台', '课程名称', '数据库课程设计'],
        ['小组成员', '郑奎昊（42411016）\n许敬淞（42411092）', '指导教师', ''],
        ['任务周次', '第二周', '任务周期', '2026-05-26 至 2026-06-01'],
        ['任务主题', '问题需求调查和数据库建模', '完成状态', '已完成'],
        ['提交日期', '2026-06-05', '文档版本', 'V1.0'],
    ]

    for i, row_data in enumerate(info_data):
        for j, text in enumerate(row_data):
            cell = info_table.cell(i, j)
            cell.text = text
            # 奇数列（标签列）加粗并设置背景色
            if j % 2 == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'D9E2F3')

    # ===== 任务清单表 =====
    doc.add_heading('二、任务完成清单', level=2)

    task_table = doc.add_table(rows=5, cols=6, style='Table Grid')
    task_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    task_headers = ['序号', '任务名称', '任务描述', '产出物', '完成状态', '完成日期']
    for j, header in enumerate(task_headers):
        cell = task_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    task_data = [
        ['1', '需求分析', '梳理业务流程，明确系统功能需求，绘制数据流图(DFD)，编写数据字典', '需求分析.md', '已完成', '2026-05-26'],
        ['2', '概念结构设计', '识别实体、属性和联系，绘制E-R图（6个实体，7个联系）', 'E-R图.md', '已完成', '2026-05-26'],
        ['3', '逻辑结构设计', '将E-R图转换为关系模式，设计6张表的结构，进行规范化分析（达到3NF）', '表结构设计.md', '已完成', '2026-05-26'],
        ['4', '物理结构设计', '编写SQL Server建库建表DDL脚本，定义约束和索引', '01_create_tables.sql', '已完成', '2026-05-26'],
    ]

    for i, row_data in enumerate(task_data):
        for j, text in enumerate(row_data):
            cell = task_table.cell(i + 1, j)
            cell.text = text
            if j == 4:  # 完成状态列
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                cell.paragraphs[0].runs[0].bold = True

    # ===== 需求分析概要表 =====
    doc.add_heading('三、需求分析概要', level=2)

    doc.add_heading('3.1 用户角色', level=3)
    role_table = doc.add_table(rows=3, cols=3, style='Table Grid')
    role_headers = ['角色', '说明', '主要操作']
    for j, header in enumerate(role_headers):
        cell = role_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    role_data = [
        ['管理员', '保卫处工作人员，系统主要操作者', '审批注册、录入违规、处理处罚、查询统计'],
        ['车主', '校园内机动车拥有者（教职工/学生/临时访客）', '提交注册申请、查询违规记录、处理处罚'],
    ]
    for i, row_data in enumerate(role_data):
        for j, text in enumerate(row_data):
            role_table.cell(i + 1, j).text = text

    doc.add_heading('3.2 功能需求', level=3)
    func_table = doc.add_table(rows=5, cols=3, style='Table Grid')
    func_headers = ['功能编号', '功能名称', '功能描述']
    for j, header in enumerate(func_headers):
        cell = func_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    func_data = [
        ['F1', '车辆权限注册管理', '车主提交注册申请，管理员审核并发放通行证，支持续期和注销'],
        ['F2', '违规记录管理', '管理员录入违规信息，系统自动关联车辆和车主，按类型扣减分数'],
        ['F3', '处罚处理', '车主查询违规记录，管理员执行处罚，累计积分达阈值自动触发处罚'],
        ['F4', '数据查询与统计', '按车牌/姓名/时间段/违规类型/学院统计违规情况'],
    ]
    for i, row_data in enumerate(func_data):
        for j, text in enumerate(row_data):
            func_table.cell(i + 1, j).text = text

    # ===== 概念设计概要表 =====
    doc.add_heading('四、概念设计概要（E-R模型）', level=2)

    doc.add_heading('4.1 实体清单', level=3)
    entity_table = doc.add_table(rows=7, cols=4, style='Table Grid')
    entity_headers = ['序号', '实体名称', '主键', '属性数量']
    for j, header in enumerate(entity_headers):
        cell = entity_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    entity_data = [
        ['1', '车主 Owner', 'owner_id', '7'],
        ['2', '车辆 Vehicle', 'plate_number', '4'],
        ['3', '通行证 Permit', 'permit_id', '7'],
        ['4', '违规记录 Violation', 'violation_id', '8'],
        ['5', '处罚记录 Penalty', 'penalty_id', '7'],
        ['6', '管理员 Admin', 'admin_id', '3'],
    ]
    for i, row_data in enumerate(entity_data):
        for j, text in enumerate(row_data):
            entity_table.cell(i + 1, j).text = text

    doc.add_heading('4.2 联系清单', level=3)
    rel_table = doc.add_table(rows=8, cols=5, style='Table Grid')
    rel_headers = ['序号', '联系名称', '实体A', '实体B', '对应关系']
    for j, header in enumerate(rel_headers):
        cell = rel_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    rel_data = [
        ['1', '拥有', '车主', '车辆', '1:N'],
        ['2', '持有', '车辆', '通行证', '1:1'],
        ['3', '产生', '车辆', '违规记录', '1:N'],
        ['4', '对应', '违规记录', '处罚记录', '1:1'],
        ['5', '审批', '管理员', '通行证', '1:N'],
        ['6', '录入', '管理员', '违规记录', '1:N'],
        ['7', '处理', '管理员', '处罚记录', '1:N'],
    ]
    for i, row_data in enumerate(rel_data):
        for j, text in enumerate(row_data):
            rel_table.cell(i + 1, j).text = text

    # ===== 逻辑设计概要表 =====
    doc.add_heading('五、逻辑设计概要（表结构）', level=2)

    doc.add_heading('5.1 关系模式总览', level=3)
    schema_table = doc.add_table(rows=7, cols=4, style='Table Grid')
    schema_headers = ['表名', '中文名', '主键', '外键数']
    for j, header in enumerate(schema_headers):
        cell = schema_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    schema_data = [
        ['Owner', '车主信息表', 'owner_id', '0'],
        ['Vehicle', '车辆信息表', 'plate_number', '1'],
        ['Permit', '通行证表', 'permit_id', '2'],
        ['Violation', '违规记录表', 'violation_id', '2'],
        ['Penalty', '处罚记录表', 'penalty_id', '2'],
        ['Admin', '管理员表', 'admin_id', '0'],
    ]
    for i, row_data in enumerate(schema_data):
        for j, text in enumerate(row_data):
            schema_table.cell(i + 1, j).text = text

    doc.add_heading('5.2 外键关系', level=3)
    fk_table = doc.add_table(rows=8, cols=5, style='Table Grid')
    fk_headers = ['外键所在表', '外键列', '引用表', '引用列', '删除行为']
    for j, header in enumerate(fk_headers):
        cell = fk_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    fk_data = [
        ['Vehicle', 'owner_id', 'Owner', 'owner_id', 'CASCADE'],
        ['Permit', 'plate_number', 'Vehicle', 'plate_number', 'CASCADE'],
        ['Permit', 'admin_id', 'Admin', 'admin_id', 'NO ACTION'],
        ['Violation', 'plate_number', 'Vehicle', 'plate_number', 'CASCADE'],
        ['Violation', 'admin_id', 'Admin', 'admin_id', 'NO ACTION'],
        ['Penalty', 'violation_id', 'Violation', 'violation_id', 'CASCADE'],
        ['Penalty', 'admin_id', 'Admin', 'admin_id', 'NO ACTION'],
    ]
    for i, row_data in enumerate(fk_data):
        for j, text in enumerate(row_data):
            fk_table.cell(i + 1, j).text = text

    doc.add_heading('5.3 规范化分析', level=3)
    norm_table = doc.add_table(rows=4, cols=3, style='Table Grid')
    norm_headers = ['范式', '检查内容', '结论']
    for j, header in enumerate(norm_headers):
        cell = norm_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    norm_data = [
        ['1NF', '所有字段原子性，不可再分', '满足'],
        ['2NF', '所有表均为单列主键，不存在部分依赖', '满足'],
        ['3NF', '逐表检查非主属性对主键的传递依赖，均不存在', '满足'],
    ]
    for i, row_data in enumerate(norm_data):
        for j, text in enumerate(row_data):
            norm_table.cell(i + 1, j).text = text

    # ===== 业务规则表 =====
    doc.add_heading('六、核心业务规则', level=2)

    rule_table = doc.add_table(rows=7, cols=3, style='Table Grid')
    rule_headers = ['规则编号', '规则名称', '规则描述']
    for j, header in enumerate(rule_headers):
        cell = rule_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    rule_data = [
        ['R1', '注册规则', '一个车主可以注册多辆车，一辆车只能对应一个车主'],
        ['R2', '通行证规则', '每辆车同时只能持有一张有效通行证'],
        ['R3', '违规积分规则', '每次违规根据类型扣减1-6分，累计扣分达到12分暂停通行证'],
        ['R4', '处罚规则', '违停罚款50元、超速罚款100元、无证入校罚款200元、占用消防通道罚款500元'],
        ['R5', '申诉规则', '车主可在处罚生效后7日内提出申诉，由管理员审核'],
        ['R6', '通行证续期', '到期前可申请续期，过期后自动变为"过期"状态'],
    ]
    for i, row_data in enumerate(rule_data):
        for j, text in enumerate(row_data):
            rule_table.cell(i + 1, j).text = text

    # ===== 产出物清单表 =====
    doc.add_heading('七、产出物清单', level=2)

    output_table = doc.add_table(rows=7, cols=4, style='Table Grid')
    output_headers = ['序号', '文件名', '文件类型', '说明']
    for j, header in enumerate(output_headers):
        cell = output_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    output_data = [
        ['1', '需求分析.md', 'Markdown', '需求说明、DFD数据流图、数据字典、业务规则'],
        ['2', 'E-R图.md', 'Markdown', '6个实体属性定义、7个联系、E-R图（文字版）'],
        ['3', '表结构设计.md', 'Markdown', '6张表结构定义、外键关系、规范化分析'],
        ['4', '01_create_tables.sql', 'SQL', '建库建表DDL脚本（含约束、索引）'],
        ['5', '02_insert_data.sql', 'SQL', '测试数据插入脚本（每表10-20条）'],
        ['6', 'run_sql.py', 'Python', 'SQL脚本执行器'],
    ]
    for i, row_data in enumerate(output_data):
        for j, text in enumerate(row_data):
            output_table.cell(i + 1, j).text = text

    # ===== 下周计划表 =====
    doc.add_heading('八、下周计划（第三周）', level=2)

    plan_table = doc.add_table(rows=6, cols=3, style='Table Grid')
    plan_headers = ['任务', '具体内容', '预计产出']
    for j, header in enumerate(plan_headers):
        cell = plan_table.cell(0, j)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    plan_data = [
        ['视图开发', '创建常用查询视图（车主违规历史、本月违规统计等）', '03_views.sql'],
        ['存储过程开发', '实现违规处理流程、注册审批流程等存储过程', '04_procedures.sql'],
        ['触发器开发', '违规积分自动更新、通行证状态自动变更', '05_triggers.sql'],
        ['CLI界面开发', 'Python命令行交互界面，实现CRUD操作', 'main.py'],
        ['集成测试', 'CLI调用数据库，验证各功能模块', '测试报告'],
    ]
    for i, row_data in enumerate(plan_data):
        for j, text in enumerate(row_data):
            plan_table.cell(i + 1, j).text = text

    # ===== 签名区 =====
    doc.add_paragraph('')
    doc.add_paragraph('')
    sign_table = doc.add_table(rows=2, cols=4, style='Table Grid')
    sign_data = [
        ['小组成员签名', '', '日期', ''],
        ['', '', '', ''],
    ]
    for i, row_data in enumerate(sign_data):
        for j, text in enumerate(row_data):
            sign_table.cell(i, j).text = text

    # 保存文档
    output_path = 'E:\\DataBase-Homework\\docs\\第二周任务完成报告.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')

if __name__ == '__main__':
    create_week2_report()
